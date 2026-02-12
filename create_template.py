"""Script to create a vulnerability report template."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Title
title = doc.add_heading('VULNERABILITY ASSESSMENT REPORT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Executive Summary
doc.add_heading('Executive Summary', 1)

summary_para = doc.add_paragraph()
summary_para.add_run('Total Vulnerabilities Found: ').bold = True
summary_para.add_run('{{TOTAL_VULNS}}')

doc.add_paragraph()
doc.add_paragraph('Risk Distribution:')
doc.add_paragraph('• Critical: {{CRITICAL_COUNT}}')
doc.add_paragraph('• High: {{HIGH_COUNT}}')
doc.add_paragraph('• Medium: {{MEDIUM_COUNT}}')
doc.add_paragraph('• Low: {{LOW_COUNT}}')
doc.add_paragraph('• Informational: {{INFO_COUNT}}')

doc.add_page_break()

# Detailed Findings
doc.add_heading('Detailed Findings', 1)

doc.add_paragraph('The following vulnerabilities were identified during the assessment:')
doc.add_paragraph()

# Create vulnerability table template
table = doc.add_table(rows=8, cols=2)
table.style = 'Light Grid Accent 1'

# Table headers and placeholders
rows_data = [
    ('Vulnerability ID', '{{VULN_ID}}'),
    ('Title', '{{TITLE}}'),
    ('Risk Level', '{{RISK_LEVEL}}'),
    ('CVSS Score', '{{CVSS_SCORE}}'),
    ('Description', '{{DESCRIPTION}}'),
    ('Affected Components', '{{AFFECTED_COMPONENTS}}'),
    ('Recommendation', '{{RECOMMENDATION}}'),
    ('Proof of Concept', '{{POC}}'),
]

for i, (label, placeholder) in enumerate(rows_data):
    row = table.rows[i]
    row.cells[0].text = label
    # Make label bold
    for paragraph in row.cells[0].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
    
    row.cells[1].text = placeholder

# Add spacing
doc.add_paragraph()
doc.add_paragraph('---' * 30)
doc.add_paragraph()

# Footer
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_para.add_run('End of Report')
footer_run.font.italic = True
footer_run.font.size = Pt(10)
footer_run.font.color.rgb = RGBColor(128, 128, 128)

# Save
doc.save('Vulnerability_Report_Template.docx')

print("✅ Template created successfully: Vulnerability_Report_Template.docx")
print()
print("📋 This template includes:")
print("   • Executive summary with vulnerability counts")
print("   • Vulnerability table with all placeholders")
print("   • Professional formatting")
print()
print("🚀 Next steps:")
print("   1. Open Vulnerability_Report_Template.docx")
print("   2. Customize styling, add logo, adjust formatting")
print("   3. Upload to Phase 2 API with your Excel file")
print()
print("⚠️  DO NOT remove the {{PLACEHOLDERS}} - they will be replaced with actual data!")
