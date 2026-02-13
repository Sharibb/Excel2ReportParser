"""Debug script to check text box detection in Word template."""

from docx import Document
from docx.oxml.ns import qn

def check_textboxes():
    """Check for text boxes in the Word template."""
    doc = Document('WAPT-Rootnik-Technical.docx')
    
    print("=" * 80)
    print("CHECKING FOR TEXT BOXES IN WORD TEMPLATE")
    print("=" * 80)
    
    # First, check entire document text for POC
    full_text = '\n'.join([para.text for para in doc.paragraphs])
    if '{{POC}}' in full_text or 'POC' in full_text:
        print("\n[INFO] Found 'POC' somewhere in document body paragraphs")
    
    # Check document body XML
    body = doc._element.body
    all_textboxes_in_doc = body.xpath('.//w:txbxContent')
    all_vml_textboxes = body.xpath('.//v:textbox//w:txbxContent')
    print(f"\n[INFO] Total text boxes in entire document: {len(all_textboxes_in_doc) + len(all_vml_textboxes)}")
    print(f"       Standard (w:txbxContent): {len(all_textboxes_in_doc)}")
    print(f"       VML (v:textbox): {len(all_vml_textboxes)}")
    
    # Check each text box in document
    for idx, txbx in enumerate(all_textboxes_in_doc + all_vml_textboxes):
        paragraphs = txbx.xpath('.//w:p')
        for p_elem in paragraphs:
            p_text = ''.join([
                t.text for t in p_elem.xpath('.//w:t')
                if t.text
            ])
            if 'POC' in p_text or '{{' in p_text:
                print(f"\n[TEXT BOX {idx + 1}] Contains: {p_text}")
    
    # Check all tables
    for table_idx, table in enumerate(doc.tables):
        print(f"\n[TABLE {table_idx + 1}]: {len(table.rows)} rows x {len(table.columns)} cols")
        
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                # Get cell XML
                tc = cell._element
                
                # Check cell text for {{POC}}
                if '{{POC}}' in cell.text:
                    print(f"\n  [FOUND] {{{{POC}}}} in Row {row_idx + 1}, Cell {cell_idx + 1}")
                    print(f"     Cell text: {cell.text[:100]}...")
                    
                    # Search for text boxes
                    textboxes = tc.xpath('.//w:txbxContent')
                    vml_textboxes = tc.xpath('.//v:textbox//w:txbxContent')
                    
                    print(f"     Standard text boxes found: {len(textboxes)}")
                    print(f"     VML text boxes found: {len(vml_textboxes)}")
                    
                    all_textboxes = textboxes + vml_textboxes
                    
                    if all_textboxes:
                        for txbx_idx, txbx in enumerate(all_textboxes):
                            print(f"\n     [TEXT BOX {txbx_idx + 1}]:")
                            
                            # Get paragraphs in text box
                            paragraphs = txbx.xpath('.//w:p')
                            print(f"        Paragraphs in text box: {len(paragraphs)}")
                            
                            for p_idx, p_elem in enumerate(paragraphs):
                                p_text = ''.join([
                                    t.text for t in p_elem.xpath('.//w:t')
                                    if t.text
                                ])
                                if p_text:
                                    print(f"        Para {p_idx + 1}: {p_text}")
                                    if '{{POC}}' in p_text:
                                        print(f"        [SUCCESS] FOUND {{{{POC}}}} IN TEXT BOX!")
                    else:
                        print(f"     [WARNING] NO TEXT BOXES FOUND - {{{{POC}}}} is directly in cell")
                        print(f"     Checking cell structure:")
                        
                        # Check for shapes
                        shapes = tc.xpath('.//w:pict')
                        print(f"     Pictures/Shapes found: {len(shapes)}")
                        
                        # Check for VML shapes
                        vml_shapes = tc.xpath('.//v:shape')
                        print(f"     VML shapes found: {len(vml_shapes)}")
                        
                        # Check for alternate content
                        alt_content = tc.xpath('.//mc:AlternateContent')
                        print(f"     Alternate content found: {len(alt_content)}")
                        
                        # Print XML structure (first 500 chars)
                        xml_str = str(tc)[:500]
                        print(f"\n     XML Preview:")
                        print(f"     {xml_str}...")

    print("\n" + "=" * 80)
    print("DIAGNOSIS COMPLETE")
    print("=" * 80)

if __name__ == "__main__":
    check_textboxes()
