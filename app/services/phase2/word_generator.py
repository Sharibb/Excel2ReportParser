"""Word document generator for populating templates with vulnerability data."""

import copy
from pathlib import Path
from typing import List, Optional, Dict, Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from PIL import Image as PILImage

from app.core.exceptions import (
    TemplateError,
    CorruptedTemplateError,
    TemplateMismatchError,
)
from app.core.logging import get_logger
from app.models.vulnerability import Vulnerability, VulnerabilityReport
from app.utils.validators import validate_docx_file

logger = get_logger(__name__)


class WordGenerator:
    """Generator for populating Word templates with vulnerability data."""

    def __init__(self, template_path: Path) -> None:
        """
        Initialize Word generator.

        Args:
            template_path: Path to Word template file
        """
        self.template_path = template_path
        self.document: Optional[Document] = None
        self.poc_base_path: Optional[Path] = None

    def generate(
        self,
        report: VulnerabilityReport,
        output_path: Path,
        poc_base_path: Optional[Path] = None,
    ) -> Path:
        """
        Generate Word document from template and vulnerability data.

        Args:
            report: VulnerabilityReport with vulnerability data
            output_path: Path where generated document should be saved
            poc_base_path: Base path for PoC image folders

        Returns:
            Path to generated document

        Raises:
            TemplateError: If template processing fails
        """
        try:
            # Validate template
            validate_docx_file(self.template_path)

            # Load document
            self.document = Document(str(self.template_path))
            self.poc_base_path = poc_base_path

            # Validate template structure
            self._validate_template_structure()

            # Sort vulnerabilities by risk level (Critical, High, Medium, Low, Informational)
            sorted_vulns = self._sort_vulnerabilities_by_risk(report.vulnerabilities)
            
            # Find ALL vulnerability table templates and process by risk level
            self._process_all_vulnerability_sections(sorted_vulns)

            # Replace placeholders in document
            self._replace_placeholders(report)

            # Save document
            output_path.parent.mkdir(parents=True, exist_ok=True)
            self.document.save(str(output_path))

            logger.info(
                f"Successfully generated Word document with {len(report.vulnerabilities)} "
                f"vulnerabilities at {output_path}"
            )

            return output_path

        except (TemplateError, CorruptedTemplateError, TemplateMismatchError):
            raise
        except Exception as e:
            logger.error(f"Failed to generate Word document: {e}")
            raise TemplateError(f"Failed to generate Word document: {str(e)}")

    def _validate_template_structure(self) -> None:
        """
        Validate that template has expected structure.

        Raises:
            CorruptedTemplateError: If template structure is invalid
        """
        if not self.document:
            raise CorruptedTemplateError("Document not loaded")

        # Basic validation: check document is not empty
        if not self.document.element.body:
            raise CorruptedTemplateError("Template document body is empty")

        logger.debug("Template structure validation passed")

    def _sort_vulnerabilities_by_risk(self, vulnerabilities: List[Vulnerability]) -> List[Vulnerability]:
        """
        Sort vulnerabilities by risk level (Critical -> High -> Medium -> Low -> Informational).

        Args:
            vulnerabilities: List of vulnerabilities to sort

        Returns:
            Sorted list of vulnerabilities
        """
        risk_order = {
            "Critical": 0,
            "High": 1,
            "Medium": 2,
            "Low": 3,
            "Informational": 4
        }
        
        return sorted(vulnerabilities, key=lambda v: risk_order.get(v.risk_level, 999))
    
    def _process_all_vulnerability_sections(self, vulnerabilities: List[Vulnerability]) -> None:
        """
        Process all vulnerability sections in the document.
        Finds all template tables and populates each with matching risk level vulnerabilities.

        Args:
            vulnerabilities: Sorted list of all vulnerabilities
        """
        if not self.document:
            return
        
        logger.info(f"Processing {len(vulnerabilities)} vulnerabilities across all risk levels")
        
        # Group vulnerabilities by risk level
        vulns_by_risk = {
            "Critical": [v for v in vulnerabilities if v.risk_level == "Critical"],
            "High": [v for v in vulnerabilities if v.risk_level == "High"],
            "Medium": [v for v in vulnerabilities if v.risk_level == "Medium"],
            "Low": [v for v in vulnerabilities if v.risk_level == "Low"],
            "Informational": [v for v in vulnerabilities if v.risk_level == "Informational"],
        }
        
        logger.info(f"Vulnerability counts: Critical={len(vulns_by_risk['Critical'])}, "
                   f"High={len(vulns_by_risk['High'])}, Medium={len(vulns_by_risk['Medium'])}, "
                   f"Low={len(vulns_by_risk['Low'])}, Info={len(vulns_by_risk['Informational'])}")
        
        # Find all vulnerability table templates
        body = self.document.element.body
        tables_to_process = []
        
        for table in self.document.tables:
            table_text = self._get_table_text(table)
            
            if any(
                placeholder in table_text
                for placeholder in [
                    "{{VULN_ID}}", "{{TITLE}}", "{{DESCRIPTION}}",
                    "{{RISK}}", "{{RISK_LEVEL}}", "{{CVSS_SCORE}}", "{{CWE_ID}}",
                ]
            ):
                # Find the position and check for preceding heading
                for idx, elem in enumerate(body):
                    if elem is table._element:
                        # Look for heading paragraph before this table
                        # Check positions -1 and -2 (heading might be 1 or 2 paragraphs before table)
                        heading_text = ""
                        heading_elem = None
                        
                        from docx.oxml.ns import qn
                        
                        for offset in [1, 2, 3]:  # Check up to 3 paragraphs before
                            if idx >= offset and body[idx - offset].tag.endswith('p'):
                                check_text = ""
                                for t_elem in body[idx - offset].iter(qn('w:t')):
                                    if t_elem.text:
                                        check_text += t_elem.text
                                
                                # Check for various heading placeholder patterns
                                # Standard: {{VULN_ID}} {{TITLE}}
                                # Malformed: {{VULN_ID}.{TITLE}}
                                if any(pattern in check_text for pattern in [
                                    '{{VULN_ID}}', '{TITLE}}', 'VULN_ID', 'TITLE'
                                ]):
                                    heading_elem = body[idx - offset]
                                    heading_text = check_text
                                    logger.info(f"Found heading at offset -{offset} with text: '{heading_text[:100]}'")
                                    break
                        
                        # Determine which section this table belongs to by looking at preceding content
                        section_risk_level = self._determine_section_risk_level(idx, body)
                        
                        tables_to_process.append({
                            'table': table,
                            'position': idx,
                            'heading': heading_elem,
                            'risk_level': section_risk_level
                        })
                        
                        logger.info(f"Found template table at position {idx} for {section_risk_level} risk section")
                        break
        
        # Process each table with its matching vulnerabilities (in reverse to handle deletions)
        for table_info in reversed(tables_to_process):
            risk_level = table_info['risk_level']
            matching_vulns = vulns_by_risk.get(risk_level, [])
            
            logger.info(f"Processing {risk_level} section with {len(matching_vulns)} vulnerabilities")
            
            if matching_vulns:
                self._generate_vulnerability_tables_for_section(
                    table_info['table'], 
                    table_info['heading'],
                    matching_vulns
                )
            else:
                # No vulnerabilities for this section - remove the template
                logger.info(f"No {risk_level} vulnerabilities - removing template table")
                table_info['table']._element.getparent().remove(table_info['table']._element)
                if table_info['heading'] is not None:
                    table_info['heading'].getparent().remove(table_info['heading'])
    
    def _determine_section_risk_level(self, table_position: int, body) -> str:
        """
        Determine which risk level section a table belongs to by examining preceding text.

        Args:
            table_position: Position of table in document body
            body: Document body element

        Returns:
            Risk level name (Critical, High, Medium, Low, or Informational)
        """
        from docx.oxml.ns import qn
        
        # Look backwards through the document to find section heading
        for idx in range(table_position - 1, -1, -1):
            if body[idx].tag.endswith('p'):
                para_text = ''
                for t_elem in body[idx].iter(qn('w:t')):
                    if t_elem.text:
                        para_text += t_elem.text
                
                para_lower = para_text.lower()
                
                if 'critical' in para_lower and 'risk' in para_lower:
                    return "Critical"
                elif 'high' in para_lower and 'risk' in para_lower:
                    return "High"
                elif 'medium' in para_lower and 'risk' in para_lower:
                    return "Medium"
                elif 'low' in para_lower and 'risk' in para_lower:
                    return "Low"
                elif 'info' in para_lower and ('finding' in para_lower or 'risk' in para_lower):
                    return "Informational"
        
        # Default to High if can't determine
        logger.warning(f"Could not determine risk level for table at position {table_position}, defaulting to High")
        return "High"
    
    def _find_vulnerability_table_template(self) -> Optional[Table]:
        """
        Find the vulnerability table template in the document.

        This looks for a table containing placeholder text like {{VULN_ID}}

        Returns:
            Table object or None if not found
        """
        if not self.document:
            return None

        for table in self.document.tables:
            # Check if table contains vulnerability placeholders
            table_text = self._get_table_text(table)

            if any(
                placeholder in table_text
                for placeholder in [
                    "{{VULN_ID}}",
                    "{{TITLE}}",
                    "{{DESCRIPTION}}",
                    "{{RISK}}",
                    "{{RISK_LEVEL}}",
                    "{{CVSS_SCORE}}",
                    "{{CWE_ID}}",
                ]
            ):
                logger.info(f"Found vulnerability table template with placeholders")
                return table

        return None

    def _get_table_text(self, table: Table) -> str:
        """
        Get all text content from a table.

        Args:
            table: Table object

        Returns:
            Concatenated text from all cells
        """
        text_parts = []
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)
        return " ".join(text_parts)

    def _generate_vulnerability_tables_for_section(
        self, template_table: Table, template_heading, vulnerabilities: List[Vulnerability]
    ) -> None:
        """
        Generate vulnerability tables for a specific risk level section.

        CRITICAL: This method MUST preserve template table structure exactly.

        Args:
            template_table: Template table to duplicate
            template_heading: Template heading element (or None)
            vulnerabilities: List of vulnerabilities to insert (already filtered by risk level)
        """
        if not self.document:
            return

        # Find the position of the template table in the document body
        body = self.document.element.body
        template_position = None
        
        for idx, elem in enumerate(body):
            if elem is template_table._element:
                template_position = idx
                # If we have a heading and it's right before the table, adjust position
                if template_heading is not None and idx > 0 and body[idx - 1] is template_heading:
                    template_position = idx - 1  # Start from heading position
                break
        
        if template_position is None:
            logger.error("Could not find template table position in document body")
            return

        logger.info(f"Found template table at position {template_position} in document body")
        if template_heading is not None:
            logger.info("Template heading provided - will duplicate for each vulnerability")
        else:
            logger.warning("No template heading provided - headings will not be generated")

        # Remove template table and heading from document (we'll replace with actual data)
        self._remove_table(template_table)
        if template_heading is not None and template_heading.getparent() is not None:
            template_heading.getparent().remove(template_heading)

        # For each vulnerability, duplicate the heading and table
        current_position = template_position
        
        for idx, vuln in enumerate(vulnerabilities):
            # If there was a template heading, duplicate it first
            if template_heading is not None:
                new_heading_element = copy.deepcopy(template_heading)
                if current_position >= len(body):
                    body.append(new_heading_element)
                else:
                    body.insert(current_position, new_heading_element)
                
                # Replace placeholders in heading - handle multiple formats including malformed ones
                # IMPORTANT: Order matters! Replace combined patterns FIRST, then individual parts
                replacements = {
                    # Malformed combined patterns (MUST BE FIRST)
                    "{{VULN_ID}.{TITLE}}": f"{vuln.vuln_id}. {vuln.title}",
                    "{{VULN_ID}}.{{TITLE}}": f"{vuln.vuln_id}. {vuln.title}",
                    "{{VULN_ID}}. {{TITLE}}": f"{vuln.vuln_id}. {vuln.title}",
                    # Individual parts (AFTER combined patterns)
                    "{{VULN_ID}}": vuln.vuln_id,
                    "{{TITLE}}": vuln.title,
                    "{TITLE}}": vuln.title,  # Malformed - single opening brace
                    # Clean up any leftover artifacts
                    "}.{": ". ",  # Remove malformed separator artifacts
                    ".{": ". ",  # Remove partial malformed separators
                }
                
                # Get heading text before replacement
                from docx.oxml.ns import qn
                heading_text_before = ""
                for t_elem in new_heading_element.iter(qn('w:t')):
                    if t_elem.text:
                        heading_text_before += t_elem.text
                
                logger.info(f"Heading before replacement: '{heading_text_before}'")
                logger.info(f"Replacing with: VULN_ID={vuln.vuln_id}, TITLE={vuln.title}")
                
                self._replace_element_text(new_heading_element, replacements)
                
                # Get heading text after replacement
                heading_text_after = ""
                for t_elem in new_heading_element.iter(qn('w:t')):
                    if t_elem.text:
                        heading_text_after += t_elem.text
                
                logger.info(f"Heading after replacement: '{heading_text_after}'")
                current_position += 1
                logger.info(f"Inserted and populated heading for {vuln.vuln_id}: {vuln.title}")
            else:
                logger.warning(f"Skipping heading for {vuln.vuln_id} - no template heading available")
            
            # Create deep copy of table XML element
            new_table_element = copy.deepcopy(template_table._element)

            # Insert new table at the current position in the body
            if current_position >= len(body):
                body.append(new_table_element)
            else:
                body.insert(current_position, new_table_element)

            # Find the newly inserted table in the document.tables collection
            inserted_table = None
            for table in self.document.tables:
                if table._element is new_table_element:
                    inserted_table = table
                    break

            if inserted_table:
                # Populate table with vulnerability data
                self._populate_vulnerability_table(inserted_table, vuln)
                logger.debug(f"Populated vulnerability table for {vuln.vuln_id}")
            else:
                logger.warning(f"Could not find inserted table for {vuln.vuln_id}")

            # Move position forward for next table
            current_position += 1
            
            # Add spacing paragraph between tables (except after the last one)
            if idx < len(vulnerabilities) - 1:
                # Create a new paragraph element for spacing
                p_element = OxmlElement('w:p')
                if current_position >= len(body):
                    body.append(p_element)
                else:
                    body.insert(current_position, p_element)
                current_position += 1
                logger.debug(f"Added spacing paragraph after {vuln.vuln_id}")

        logger.info(f"Generated {len(vulnerabilities)} vulnerability tables")

    def _get_table_index(self, table: Table) -> Optional[int]:
        """
        Get index of table in document.

        Args:
            table: Table object

        Returns:
            Index or None if not found
        """
        if not self.document:
            return None

        try:
            # Find table by comparing XML elements instead of object references
            for idx, doc_table in enumerate(self.document.tables):
                if doc_table._element is table._element:
                    return idx
            
            logger.error("Table element not found in document tables")
            return None
        except Exception as e:
            logger.error(f"Error finding table index: {e}")
            return None

    def _remove_table(self, table: Table) -> None:
        """
        Remove table from document.

        Args:
            table: Table to remove
        """
        if table._element.getparent() is not None:
            table._element.getparent().remove(table._element)

    def _insert_table_element(self, table_element: OxmlElement, index: int) -> None:
        """
        Insert table element at specific index in document body.

        Args:
            table_element: Table XML element
            index: Position to insert
        """
        if not self.document:
            return

        body = self.document.element.body

        # Find all block-level elements (paragraphs and tables)
        block_elements = [
            elem
            for elem in body
            if elem.tag.endswith(("p", "tbl"))
        ]

        # Insert at the correct position
        if index >= len(block_elements):
            body.append(table_element)
        else:
            body.insert(body.index(block_elements[index]), table_element)

    def _add_spacing_paragraph(self, index: int) -> None:
        """
        Add empty paragraph for spacing.

        Args:
            index: Position to insert paragraph
        """
        if not self.document:
            return

        # Create paragraph element
        p = OxmlElement("w:p")

        # Insert into body
        body = self.document.element.body
        block_elements = [elem for elem in body if elem.tag.endswith(("p", "tbl"))]

        if index >= len(block_elements):
            body.append(p)
        else:
            body.insert(body.index(block_elements[index]), p)

    def _populate_vulnerability_table(
        self, table: Table, vuln: Vulnerability
    ) -> None:
        """
        Populate table with vulnerability data by replacing placeholders.

        Args:
            table: Table to populate
            vuln: Vulnerability data
        """
        # Define placeholder mappings
        replacements = {
            "{{VULN_ID}}": vuln.vuln_id,
            "{{TITLE}}": vuln.title,
            "{{DESCRIPTION}}": vuln.description,
            "{{RISK}}": vuln.risk_level,
            "{{RISK_LEVEL}}": vuln.risk_level,
            "{{CVSS}}": str(vuln.cvss_score) if vuln.cvss_score else "N/A",
            "{{CVSS_SCORE}}": str(vuln.cvss_score) if vuln.cvss_score else "N/A",
            "{{AFFECTED}}": vuln.affected_components,
            "{{AFFECTED_COMPONENTS}}": vuln.affected_components,
            "{{RECOMMENDATION}}": vuln.recommendation,
            "{{REMEDIATION}}": vuln.recommendation,
            "{{CWE_ID}}": vuln.cwe_id if vuln.cwe_id else "N/A",
            "{{IMPACT}}": vuln.impact if vuln.impact else "N/A",
            "{{REFERENCES}}": vuln.references if vuln.references else "N/A",
            "{{REMEDIATION_EFFORT}}": vuln.remediation_effort if vuln.remediation_effort else "N/A",
        }

        # Replace placeholders in all cells
        for row in table.rows:
            for cell in row.cells:
                self._replace_cell_text(cell, replacements)

                # Check if this cell should contain PoC images
                if "{{POC}}" in cell.text or "{{STEPS}}" in cell.text:
                    # Check if there's a text box with {{POC}} placeholder
                    textbox_found = self._insert_poc_in_textbox(cell, vuln)
                    
                    if not textbox_found:
                        # Fallback: If no text box found, use old method
                        # Always clear the placeholders first
                        self._clear_poc_placeholders(cell)
                        # Then insert images if available
                        self._insert_poc_images(cell, vuln)
                    else:
                        # Text box found - just add step text outside the box
                        self._clear_poc_placeholders(cell)
                        self._insert_step_text_only(cell, vuln)

    def _replace_cell_text(self, cell: _Cell, replacements: Dict[str, str]) -> None:
        """
        Replace placeholders in cell text.

        Args:
            cell: Table cell
            replacements: Dictionary of placeholder -> value mappings
        """
        # Check if cell contains special placeholders that need full replacement
        cell_text = cell.text
        needs_full_replacement = False
        replacement_key = None
        replacement_value = None
        
        for key, value in replacements.items():
            if key in cell_text:
                # For RECOMMENDATION and IMPACT, clear cell completely and insert only the actual data
                if key in ["{{RECOMMENDATION}}", "{{IMPACT}}"]:
                    needs_full_replacement = True
                    replacement_key = key
                    replacement_value = value
                    break
        
        if needs_full_replacement:
            # Clear all existing content
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.text = ""
            
            # Remove all paragraphs except the first one
            while len(cell.paragraphs) > 1:
                p = cell.paragraphs[-1]._element
                p.getparent().remove(p)
            
            # Insert the actual value as plain text
            if cell.paragraphs:
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run(replacement_value)
                run.font.name = 'Tahoma'
            else:
                # No paragraphs, create one
                p = cell.add_paragraph()
                run = p.add_run(replacement_value)
                run.font.name = 'Tahoma'
            
            logger.debug(f"Cleared cell and inserted {replacement_key} data")
            return
        
        # Normal replacement for other placeholders
        for paragraph in cell.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    # Replace in runs to preserve formatting
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
                            # Set font to Tahoma
                            run.font.name = 'Tahoma'
                    
                    # Special handling for References - make them hyperlinks
                    if key == "{{REFERENCES}}" and value and value != "N/A":
                        self._add_hyperlinks_to_cell(cell, value)
    
    def _replace_element_text(self, element: OxmlElement, replacements: Dict[str, str]) -> None:
        """
        Replace placeholders in an XML element.
        Handles cases where placeholders might be split across multiple text elements.

        Args:
            element: XML element
            replacements: Dictionary of placeholder -> value mappings
        """
        from docx.oxml.ns import qn
        
        # First pass: try to replace within individual text elements
        for t_elem in element.iter(qn('w:t')):
            if t_elem.text:
                for key, value in replacements.items():
                    if key in t_elem.text:
                        t_elem.text = t_elem.text.replace(key, value)
        
        # Second pass: handle split placeholders by reconstructing full text
        # Get all text elements
        text_elements = list(element.iter(qn('w:t')))
        if len(text_elements) > 1:
            # Build full text
            full_text = "".join([t.text or "" for t in text_elements])
            
            # Check if any placeholders exist in the full text
            replacements_found = False
            for key, value in replacements.items():
                if key in full_text:
                    # Placeholder is split across elements - need to reconstruct
                    logger.info(f"Placeholder '{key}' split across {len(text_elements)} text elements, replacing with '{value}'")
                    
                    # Replace in full text
                    full_text = full_text.replace(key, value)
                    replacements_found = True
            
            # If we found any replacements, update the text elements
            if replacements_found:
                # Clear all text elements except the first one
                text_elements[0].text = full_text
                for t_elem in text_elements[1:]:
                    t_elem.text = ""
                logger.info(f"Consolidated {len(text_elements)} text elements into one")
    
    def _add_hyperlinks_to_cell(self, cell: _Cell, references_text: str) -> None:
        """
        Convert URLs in references text to clickable hyperlinks.

        Args:
            cell: Table cell containing references
            references_text: Text containing URLs (newline-separated)
        """
        import re
        
        # Clear existing content
        cell.text = ""
        
        # Split by newlines to handle multiple references
        lines = references_text.split('\n')
        
        for idx, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Check if line contains a URL
            url_pattern = r'https?://[^\s]+'
            urls = re.findall(url_pattern, line)
            
            if urls:
                # Line contains URLs - make them clickable
                for url in urls:
                    # Add URL as hyperlink
                    paragraph = cell.add_paragraph() if idx > 0 or cell.paragraphs[0].text else cell.paragraphs[0]
                    self._add_hyperlink(paragraph, url, url)
            else:
                # Plain text line
                paragraph = cell.add_paragraph() if idx > 0 else cell.paragraphs[0]
                run = paragraph.add_run(line)
                run.font.name = 'Tahoma'
    
    def _add_hyperlink(self, paragraph: Paragraph, url: str, text: str) -> None:
        """
        Add a hyperlink to a paragraph.

        Args:
            paragraph: Paragraph to add hyperlink to
            url: URL to link to
            text: Display text for the hyperlink
        """
        from docx.oxml.shared import OxmlElement, qn
        
        # Create the w:hyperlink tag and add needed values
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), '')
        
        # Create a new run and set the text
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # Set font to Tahoma
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Tahoma')
        rFonts.set(qn('w:hAnsi'), 'Tahoma')
        rPr.append(rFonts)
        
        # Make it look like a hyperlink (blue and underlined)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        new_run.append(rPr)
        
        # Add text
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)
        
        hyperlink.append(new_run)
        
        # Add hyperlink to paragraph
        paragraph._element.append(hyperlink)
        
        # Add relationship to the document
        try:
            part = paragraph.part
            r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
            hyperlink.set(qn('r:id'), r_id)
        except Exception as e:
            logger.warning(f"Could not create hyperlink relationship: {e}")

    def _clear_poc_placeholders(self, cell: _Cell) -> None:
        """
        Clear PoC placeholder text from cell.

        Args:
            cell: Table cell to clear placeholders from
        """
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = run.text.replace("{{POC}}", "").replace("{{STEPS}}", "")

    def _insert_poc_images(self, cell: _Cell, vuln: Vulnerability) -> None:
        """
        Insert PoC steps as text and try to insert corresponding images.

        Args:
            cell: Table cell to insert PoC content into
            vuln: Vulnerability with PoC data
        """
        # If no steps, nothing to do
        if not vuln.steps:
            logger.debug(f"No PoC steps for {vuln.vuln_id}")
            return

        logger.info(f"Processing PoC for {vuln.vuln_id}: {len(vuln.steps)} steps, poc_folder={vuln.poc_folder}")

        # Always add step descriptions as text
        for idx, step_text in enumerate(vuln.steps, start=1):
            # Add step text
            paragraph = cell.add_paragraph()
            run = paragraph.add_run(f"Step {idx}: {step_text}")
            run.bold = True
            
            logger.debug(f"Added step text: Step {idx}: {step_text}")

            # Try to insert image if poc_folder and poc_base_path are provided
            if vuln.poc_folder and self.poc_base_path:
                # Image filename is just the step number: 1.png, 2.png, 3.png, etc.
                image_filename = f"{idx}.png"
                poc_folder_path = self.poc_base_path / vuln.poc_folder
                image_path = poc_folder_path / image_filename

                logger.debug(f"Looking for image at: {image_path}")

                if image_path.exists():
                    try:
                        # Add paragraph for image
                        img_paragraph = cell.add_paragraph()

                        # Use fixed dimensions to prevent table overflow: 11.56 cm x 6.92 cm
                        target_width_cm = 11.56
                        target_height_cm = 6.92
                        width_inches = target_width_cm * 0.393701  # ~4.55 inches
                        height_inches = target_height_cm * 0.393701  # ~2.72 inches

                        # Insert image with fixed dimensions
                        img_run = img_paragraph.add_run()
                        img_run.add_picture(str(image_path), width=Inches(width_inches), height=Inches(height_inches))

                        logger.info(f"✓ Inserted PoC image: {image_path} (11.56cm x 6.92cm)")

                    except Exception as e:
                        logger.warning(f"Failed to insert image {image_path}: {e}")
                else:
                    logger.debug(f"Image not found: {image_path}")
            else:
                if not vuln.poc_folder:
                    logger.debug(f"No poc_folder specified for {vuln.vuln_id}")
                if not self.poc_base_path:
                    logger.debug(f"No poc_base_path provided")

            # Add spacing between steps
            if idx < len(vuln.steps):
                cell.add_paragraph()  # Empty line between steps

    def _insert_step_text_only(self, cell: _Cell, vuln: Vulnerability) -> None:
        """
        Insert only step text (without images) outside the text box.
        
        Args:
            cell: Table cell to insert step text into
            vuln: Vulnerability with PoC data
        """
        if not vuln.steps:
            return
        
        logger.info(f"Inserting step text only for {vuln.vuln_id}")
        
        # Add step descriptions as text (images will be in text box)
        for idx, step_text in enumerate(vuln.steps, start=1):
            paragraph = cell.add_paragraph()
            run = paragraph.add_run(f"Step {idx}: {step_text}")
            run.bold = True
            
            # Add spacing between steps
            if idx < len(vuln.steps):
                cell.add_paragraph()

    def _insert_poc_in_textbox(self, cell: _Cell, vuln: Vulnerability) -> bool:
        """
        Find text box with {{POC}} placeholder and insert images inside it.
        
        Args:
            cell: Table cell that may contain text box
            vuln: Vulnerability with PoC data
            
        Returns:
            True if text box was found and images inserted, False otherwise
        """
        if not vuln.steps or not vuln.poc_folder or not self.poc_base_path:
            return False
        
        try:
            # Get the cell's underlying XML element
            tc = cell._element
            
            # Search for all txbxContent elements (text box content)
            # These can be in various locations in the XML structure
            try:
                textboxes = tc.xpath('.//w:txbxContent')
                logger.debug(f"Found {len(textboxes)} standard text boxes")
            except Exception as e:
                logger.warning(f"Error searching for standard text boxes: {e}")
                textboxes = []
            
            # Also check for VML text boxes
            try:
                vml_textboxes = tc.xpath('.//v:textbox//w:txbxContent')
                logger.debug(f"Found {len(vml_textboxes)} VML text boxes")
                textboxes.extend(vml_textboxes)
            except Exception as e:
                logger.warning(f"Error searching for VML text boxes: {e}")
            
            if not textboxes:
                logger.warning(f"No text boxes found in cell for {vuln.vuln_id}")
                return False
            
            logger.info(f"Total text boxes found: {len(textboxes)} for {vuln.vuln_id}")
            
            # Check each text box for {{POC}} placeholder
            for txbx_idx, txbx_content in enumerate(textboxes):
                logger.debug(f"Checking text box {txbx_idx + 1} for {{{{POC}}}} placeholder")
                # Get all paragraphs in the text box
                try:
                    paragraphs = txbx_content.xpath('.//w:p')
                except Exception as e:
                    logger.warning(f"Error getting paragraphs from text box: {e}")
                    continue
                
                # Check if any paragraph contains {{POC}}
                contains_poc = False
                poc_paragraph_idx = -1
                
                for idx, p_elem in enumerate(paragraphs):
                    try:
                        p_text = ''.join([
                            t.text for t in p_elem.xpath('.//w:t')
                            if t.text
                        ])
                        logger.debug(f"Text box paragraph {idx}: '{p_text}'")
                        if '{{POC}}' in p_text:
                            contains_poc = True
                            poc_paragraph_idx = idx
                            logger.info(f"Found {{{{POC}}}} in paragraph {idx}")
                            break
                    except Exception as e:
                        logger.warning(f"Error reading paragraph {idx}: {e}")
                        continue
                
                if contains_poc:
                    logger.info(f"Found text box with {{{{POC}}}} placeholder for {vuln.vuln_id}")
                    
                    # Clear the {{POC}} placeholder paragraph
                    if poc_paragraph_idx >= 0 and poc_paragraph_idx < len(paragraphs):
                        poc_para = paragraphs[poc_paragraph_idx]
                        # Remove all runs from the placeholder paragraph
                        try:
                            for run in poc_para.xpath('.//w:r'):
                                poc_para.remove(run)
                            logger.debug(f"Cleared {{{{POC}}}} placeholder from paragraph {poc_paragraph_idx}")
                        except Exception as e:
                            logger.warning(f"Error clearing placeholder: {e}")
                    
                    # Insert images into the text box
                    self._insert_images_in_textbox(txbx_content, vuln, poc_paragraph_idx)
                    return True
            
            return False
            
        except Exception as e:
            logger.warning(f"Failed to process text box for {vuln.vuln_id}: {e}")
            logger.exception(e)
            return False

    def _insert_images_in_textbox(self, txbx_content: OxmlElement, vuln: Vulnerability, insert_after_idx: int = -1) -> None:
        """
        Insert PoC images into a text box element.
        
        Args:
            txbx_content: The w:txbxContent XML element
            vuln: Vulnerability with PoC data
            insert_after_idx: Index of paragraph after which to insert images (-1 to append)
        """
        try:
            # Get existing paragraphs
            existing_paragraphs = txbx_content.findall(qn('w:p'), txbx_content.nsmap)
            
            # For each step, insert the corresponding image
            for idx, step_text in enumerate(vuln.steps, start=1):
                if vuln.poc_folder and self.poc_base_path:
                    image_filename = f"{idx}.png"
                    poc_folder_path = self.poc_base_path / vuln.poc_folder
                    image_path = poc_folder_path / image_filename
                    
                    if image_path.exists():
                        logger.debug(f"Inserting image {image_path} into text box")
                        
                        # Use fixed dimensions: 11.56 cm width x 6.92 cm height
                        # Convert cm to inches: 1 cm = 0.393701 inches
                        target_width_cm = 11.56
                        target_height_cm = 6.92
                        width_inches = target_width_cm * 0.393701  # ~4.55 inches
                        height_inches = target_height_cm * 0.393701  # ~2.72 inches
                        
                        # Create new paragraph element for the image
                        p_elem = OxmlElement('w:p')
                        
                        # Create a Paragraph object wrapper to use add_picture
                        temp_para = Paragraph(p_elem, self.document)
                        
                        # Add the image with fixed dimensions
                        run = temp_para.add_run()
                        run.add_picture(str(image_path), width=Inches(width_inches), height=Inches(height_inches))
                        
                        # Append the paragraph to the text box
                        txbx_content.append(p_elem)
                        
                        logger.info(f"✓ Inserted image {image_path} into text box (11.56cm x 6.92cm)")
                        
                        # Add spacing paragraph between images
                        if idx < len(vuln.steps):
                            spacing_p = OxmlElement('w:p')
                            # Add empty run to make the paragraph valid
                            spacing_r = OxmlElement('w:r')
                            spacing_p.append(spacing_r)
                            txbx_content.append(spacing_p)
                    else:
                        logger.debug(f"Image not found: {image_path}")
        
        except Exception as e:
            logger.warning(f"Failed to insert images in text box: {e}")
            logger.exception(e)

    def _calculate_image_width(
        self, image_path: Path, max_width: float = 6.0
    ) -> float:
        """
        Calculate appropriate width for image.

        Args:
            image_path: Path to image file
            max_width: Maximum width in inches

        Returns:
            Width in inches
        """
        try:
            with PILImage.open(image_path) as img:
                width_px, height_px = img.size

                # Calculate aspect ratio
                aspect_ratio = width_px / height_px

                # Use max width if image is wide
                if aspect_ratio > 1.5:
                    return max_width
                else:
                    return min(max_width * 0.7, max_width)

        except Exception:
            # Default width if image cannot be analyzed
            return 4.5

    def _replace_placeholders(self, report: VulnerabilityReport) -> None:
        """
        Replace global placeholders in document.

        Args:
            report: VulnerabilityReport with data
        """
        if not self.document:
            return

        # Generate findings lists for summary table
        high_findings = [v for v in report.vulnerabilities if v.risk_level == "High"]
        medium_findings = [v for v in report.vulnerabilities if v.risk_level == "Medium"]
        low_findings = [v for v in report.vulnerabilities if v.risk_level == "Low"]
        critical_findings = [v for v in report.vulnerabilities if v.risk_level == "Critical"]
        info_findings = [v for v in report.vulnerabilities if v.risk_level == "Informational"]

        replacements = {
            "{{TOTAL_VULNS}}": str(report.total_count),
            "{{CRITICAL_COUNT}}": str(report.critical_count),
            "{{HIGH_COUNT}}": str(report.high_count),
            "{{MEDIUM_COUNT}}": str(report.medium_count),
            "{{LOW_COUNT}}": str(report.low_count),
            "{{INFO_COUNT}}": str(report.info_count),
            "{{REPORT_TITLE}}": report.report_title,
            "{{APP_URL}}": "https://example.com",  # Default, can be made configurable
        }

        logger.info(f"Replacing placeholders with counts: Critical={report.critical_count}, "
                   f"High={report.high_count}, Medium={report.medium_count}, "
                   f"Low={report.low_count}, Info={report.info_count}")

        # Replace in ALL document elements using XML traversal
        from docx.oxml.ns import qn
        body = self.document.element.body
        
        # First pass: Replace in individual text elements
        for elem in body.iter():
            if elem.tag.endswith('t'):  # w:t elements contain text
                if elem.text:
                    original_text = elem.text
                    for key, value in replacements.items():
                        if key in elem.text:
                            elem.text = elem.text.replace(key, value)
                            if original_text != elem.text:
                                logger.info(f"Replaced '{key}' with '{value}' in text element")
        
        # Second pass: Handle split placeholders in paragraphs
        # This handles cases where {{MEDIUM_COUNT}} is split as "{{" + "MEDIUM" + "_COUNT}}"
        for paragraph_elem in body.iter():
            if paragraph_elem.tag.endswith('p'):  # w:p elements are paragraphs
                # Get all text elements in this paragraph
                text_elements = list(paragraph_elem.iter(qn('w:t')))
                if len(text_elements) > 1:
                    # Build full paragraph text
                    full_text = "".join([t.text or "" for t in text_elements])
                    
                    # Check if any placeholders exist in the full text
                    replacements_found = False
                    original_full_text = full_text
                    for key, value in replacements.items():
                        if key in full_text:
                            logger.info(f"Found split placeholder '{key}' in paragraph, replacing with '{value}'")
                            full_text = full_text.replace(key, value)
                            replacements_found = True
                    
                    # If we found any replacements, update the text elements
                    if replacements_found:
                        logger.info(f"Consolidated paragraph text from '{original_full_text[:100]}' to '{full_text[:100]}'")
                        # Put all text in the first element, clear the rest
                        text_elements[0].text = full_text
                        for t_elem in text_elements[1:]:
                            t_elem.text = ""

        # Process summary table to insert findings as separate rows
        self._populate_summary_table(report, high_findings, medium_findings, low_findings, critical_findings, info_findings)

        logger.debug("Replaced global placeholders")
    
    def _populate_summary_table(
        self, 
        report: VulnerabilityReport,
        high_findings: List[Vulnerability],
        medium_findings: List[Vulnerability],
        low_findings: List[Vulnerability],
        critical_findings: List[Vulnerability],
        info_findings: List[Vulnerability]
    ) -> None:
        """
        Populate summary table with findings, creating separate rows for each finding.

        Args:
            report: VulnerabilityReport with data
            high_findings: List of high-risk vulnerabilities
            medium_findings: List of medium-risk vulnerabilities
            low_findings: List of low-risk vulnerabilities
            critical_findings: List of critical-risk vulnerabilities
            info_findings: List of informational findings
        """
        if not self.document:
            return

        findings_map = {
            "{{HIGH_FINDINGS_LIST}}": (high_findings, "HIGH"),
            "{{MEDIUM_FINDINGS_LIST}}": (medium_findings, "MEDIUM"),
            "{{LOW_FINDINGS_LIST}}": (low_findings, "LOW"),
            "{{CRITICAL_FINDINGS_LIST}}": (critical_findings, "CRITICAL"),
            "{{INFO_FINDINGS_LIST}}": (info_findings, "INFO"),
        }

        # Find and process summary table
        for table_idx, table in enumerate(self.document.tables):
            rows_to_process = []
            
            # Find rows with placeholder markers
            for row_idx, row in enumerate(table.rows):
                row_text = "".join([cell.text for cell in row.cells])
                
                for placeholder, (findings, status) in findings_map.items():
                    if placeholder in row_text:
                        logger.info(f"Found placeholder '{placeholder}' in table {table_idx}, row {row_idx}: '{row_text[:100]}'")
                        rows_to_process.append((row_idx, placeholder, findings, status, row))
                        break
            
            if rows_to_process:
                logger.info(f"Table {table_idx} has {len(rows_to_process)} rows to process")
            
            # Process each placeholder row (in reverse to handle insertions correctly)
            for row_idx, placeholder, findings, status, template_row in reversed(rows_to_process):
                if findings:
                    # Get the position where we'll insert rows
                    template_row_element = template_row._element
                    parent_table = table._element
                    template_position = parent_table.index(template_row_element)
                    
                    # Insert rows for each finding (in order)
                    for finding_idx, vuln in enumerate(findings):
                        # Deep copy the template row XML
                        new_row_element = copy.deepcopy(template_row_element)
                        
                        # Insert the new row right after the template row position
                        insert_position = template_position + 1 + finding_idx
                        parent_table.insert(insert_position, new_row_element)
                        
                        # Now find the newly created row in the table.rows collection and populate it
                        # The row should be at the position we just inserted
                        new_row = None
                        for tbl_row in table.rows:
                            if tbl_row._element is new_row_element:
                                new_row = tbl_row
                                break
                        
                        if new_row:
                            # Replace placeholder in first cell with actual finding
                            # Use proper split-placeholder handling
                            cell0 = new_row.cells[0]
                            cell0_replaced = False
                            
                            for paragraph in cell0.paragraphs:
                                # Try simple replacement first
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, f"{vuln.vuln_id}. {vuln.title}")
                                        cell0_replaced = True
                                
                                # If not replaced, handle split placeholders
                                if not cell0_replaced and paragraph.runs:
                                    # Build full text from all runs
                                    full_text = "".join([run.text or "" for run in paragraph.runs])
                                    
                                    if placeholder in full_text:
                                        # Placeholder is split - consolidate
                                        logger.info(f"Split placeholder '{placeholder}' detected in summary table, consolidating")
                                        full_text = full_text.replace(placeholder, f"{vuln.vuln_id}. {vuln.title}")
                                        
                                        # Put all text in first run, clear the rest
                                        paragraph.runs[0].text = full_text
                                        for run in paragraph.runs[1:]:
                                            run.text = ""
                                        cell0_replaced = True
                            
                            # Status cell should already have the correct formatting from template
                            # Just verify it has content
                            cell1 = new_row.cells[1]
                            if not cell1.text.strip() or placeholder in cell1.text:
                                # Set status if not already set
                                for paragraph in cell1.paragraphs:
                                    if paragraph.runs:
                                        paragraph.runs[0].text = status
                                    else:
                                        run = paragraph.add_run(status)
                                        run.bold = True
                    
                    # Remove the original template row
                    template_row_element.getparent().remove(template_row_element)
                    logger.info(f"Inserted {len(findings)} rows for {placeholder}")
                else:
                    # No findings - replace placeholder with "None"
                    # Handle split placeholders properly
                    for cell in template_row.cells:
                        cell_replaced = False
                        for paragraph in cell.paragraphs:
                            # Try simple replacement first
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, "None")
                                    cell_replaced = True
                            
                            # If not replaced, handle split placeholders
                            if not cell_replaced and paragraph.runs:
                                # Build full text from all runs
                                full_text = "".join([run.text or "" for run in paragraph.runs])
                                
                                if placeholder in full_text:
                                    # Placeholder is split - consolidate
                                    logger.info(f"Split placeholder '{placeholder}' detected (no findings case), replacing with 'None'")
                                    full_text = full_text.replace(placeholder, "None")
                                    
                                    # Put all text in first run, clear the rest
                                    paragraph.runs[0].text = full_text
                                    for run in paragraph.runs[1:]:
                                        run.text = ""
                                    cell_replaced = True
                    
                    logger.info(f"No findings for {placeholder}, marked as None")

        logger.debug("Populated summary table with individual rows for each finding")
