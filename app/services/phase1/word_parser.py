"""Word document parser for extracting vulnerability data."""

import re
from pathlib import Path
from typing import List, Optional, Dict, Any

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.core.exceptions import ParsingError
from app.core.logging import get_logger
from app.models.vulnerability import Vulnerability, VulnerabilityReport, RiskLevel
from app.utils.validators import validate_docx_file, validate_vulnerability_id

logger = get_logger(__name__)


class WordParser:
    """Parser for extracting vulnerability data from Word documents."""

    def __init__(self, docx_path: Path) -> None:
        """
        Initialize Word parser.

        Args:
            docx_path: Path to Word document
        """
        self.docx_path = docx_path
        self.document: Optional[Document] = None

    def parse(self) -> VulnerabilityReport:
        """
        Parse Word document and extract vulnerabilities.

        Returns:
            VulnerabilityReport containing extracted data

        Raises:
            ParsingError: If parsing fails
        """
        try:
            # Validate file
            validate_docx_file(self.docx_path)

            # Load document
            self.document = Document(str(self.docx_path))

            # Extract report title
            report_title = self._extract_report_title()

            # Extract vulnerabilities
            vulnerabilities = self._extract_vulnerabilities()

            # Create report
            report = VulnerabilityReport(
                report_title=report_title,
                vulnerabilities=vulnerabilities,
            )
            report.calculate_counts()

            logger.info(
                f"Successfully parsed {len(vulnerabilities)} vulnerabilities "
                f"from {self.docx_path}"
            )

            return report

        except Exception as e:
            logger.error(f"Failed to parse Word document: {e}")
            raise ParsingError(f"Failed to parse Word document: {str(e)}")

    def _extract_report_title(self) -> str:
        """
        Extract report title from document.

        Returns:
            Report title or default value
        """
        if not self.document:
            return "Vulnerability Report"

        # Try to get title from first heading or paragraph
        for paragraph in self.document.paragraphs[:5]:
            if paragraph.text.strip():
                return paragraph.text.strip()

        return "Vulnerability Report"

    def _extract_vulnerabilities(self) -> List[Vulnerability]:
        """
        Extract all vulnerabilities from document.

        Returns:
            List of Vulnerability objects
        """
        if not self.document:
            return []

        vulnerabilities: List[Vulnerability] = []

        # Strategy 1: Extract from tables
        table_vulns = self._extract_from_tables()
        vulnerabilities.extend(table_vulns)

        # Strategy 2: Extract from structured sections
        section_vulns = self._extract_from_sections()
        vulnerabilities.extend(section_vulns)

        # Remove duplicates based on vuln_id
        unique_vulns = self._deduplicate_vulnerabilities(vulnerabilities)

        return unique_vulns

    def _extract_from_tables(self) -> List[Vulnerability]:
        """
        Extract vulnerabilities from tables in the document.

        Returns:
            List of vulnerabilities found in tables
        """
        vulnerabilities: List[Vulnerability] = []

        if not self.document:
            return vulnerabilities

        for table in self.document.tables:
            try:
                vuln = self._parse_vulnerability_table(table)
                if vuln:
                    vulnerabilities.append(vuln)
            except Exception as e:
                logger.warning(f"Failed to parse table: {e}")
                continue

        logger.info(f"Extracted {len(vulnerabilities)} vulnerabilities from tables")
        return vulnerabilities

    def _parse_vulnerability_table(self, table: Table) -> Optional[Vulnerability]:
        """
        Parse a single vulnerability table.

        Args:
            table: Word table object

        Returns:
            Vulnerability object or None if parsing fails
        """
        try:
            # Extract key-value pairs from table
            data: Dict[str, str] = {}

            for row in table.rows:
                cells = row.cells
                if len(cells) >= 2:
                    key = cells[0].text.strip().lower()
                    value = cells[1].text.strip()
                    data[key] = value

            # Map table fields to vulnerability fields
            vuln_id = self._find_field(
                data, ["id", "vulnerability id", "vuln id", "identifier"]
            )
            title = self._find_field(
                data, ["title", "name", "vulnerability name", "vulnerability"]
            )
            description = self._find_field(
                data, ["description", "details", "summary", "overview"]
            )
            risk_level = self._find_field(
                data, ["risk", "risk level", "severity", "impact"]
            )
            cvss_score_str = self._find_field(
                data, ["cvss", "cvss score", "score"]
            )
            affected = self._find_field(
                data,
                ["affected", "affected components", "affected systems", "component"],
            )
            recommendation = self._find_field(
                data,
                ["recommendation", "remediation", "fix", "solution", "mitigation"],
            )

            # Validate required fields
            if not (vuln_id and title and description and risk_level):
                return None

            # Validate vulnerability ID format
            try:
                vuln_id = validate_vulnerability_id(vuln_id)
            except Exception:
                return None

            # Parse CVSS score
            cvss_score = self._parse_cvss_score(cvss_score_str)

            # Map risk level
            risk_enum = self._map_risk_level(risk_level)

            return Vulnerability(
                vuln_id=vuln_id,
                title=title,
                description=description,
                risk_level=risk_enum,
                cvss_score=cvss_score,
                affected_components=affected or "N/A",
                recommendation=recommendation or "N/A",
            )

        except Exception as e:
            logger.debug(f"Could not parse table as vulnerability: {e}")
            return None

    def _extract_from_sections(self) -> List[Vulnerability]:
        """
        Extract vulnerabilities from structured text sections.

        Returns:
            List of vulnerabilities found in sections
        """
        vulnerabilities: List[Vulnerability] = []

        if not self.document:
            return vulnerabilities

        # Look for patterns like "H1: Title", "M1. Title", etc.
        pattern = r"^([CHMLI]\d+)[:\.\s]+(.*?)$"

        current_vuln_data: Optional[Dict[str, Any]] = None

        for paragraph in self.document.paragraphs:
            text = paragraph.text.strip()

            if not text:
                continue

            # Check if this starts a new vulnerability
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                # Save previous vulnerability if exists
                if current_vuln_data:
                    vuln = self._build_vulnerability_from_data(current_vuln_data)
                    if vuln:
                        vulnerabilities.append(vuln)

                # Start new vulnerability
                vuln_id = match.group(1).upper()
                title = match.group(2).strip()

                current_vuln_data = {
                    "vuln_id": vuln_id,
                    "title": title,
                    "description": "",
                    "risk_level": "",
                    "affected": "",
                    "recommendation": "",
                }
            elif current_vuln_data:
                # Accumulate text for current vulnerability
                # Try to identify field types
                text_lower = text.lower()

                if any(kw in text_lower for kw in ["description:", "details:"]):
                    current_vuln_data["description"] += text + "\n"
                elif any(kw in text_lower for kw in ["risk:", "severity:"]):
                    current_vuln_data["risk_level"] += text + "\n"
                elif any(kw in text_lower for kw in ["affected:", "component:"]):
                    current_vuln_data["affected"] += text + "\n"
                elif any(
                    kw in text_lower
                    for kw in ["recommendation:", "remediation:", "fix:"]
                ):
                    current_vuln_data["recommendation"] += text + "\n"
                else:
                    # Default: add to description
                    current_vuln_data["description"] += text + "\n"

        # Don't forget the last vulnerability
        if current_vuln_data:
            vuln = self._build_vulnerability_from_data(current_vuln_data)
            if vuln:
                vulnerabilities.append(vuln)

        logger.info(f"Extracted {len(vulnerabilities)} vulnerabilities from sections")
        return vulnerabilities

    def _build_vulnerability_from_data(
        self, data: Dict[str, Any]
    ) -> Optional[Vulnerability]:
        """
        Build Vulnerability object from extracted data dictionary.

        Args:
            data: Dictionary with vulnerability fields

        Returns:
            Vulnerability object or None
        """
        try:
            vuln_id = data.get("vuln_id", "").strip()
            title = data.get("title", "").strip()
            description = data.get("description", "").strip()
            risk_text = data.get("risk_level", "").strip()
            affected = data.get("affected", "").strip()
            recommendation = data.get("recommendation", "").strip()

            if not (vuln_id and title):
                return None

            # Infer risk level from vuln_id if not provided
            if not risk_text:
                risk_level = self._infer_risk_from_id(vuln_id)
            else:
                risk_level = self._map_risk_level(risk_text)

            return Vulnerability(
                vuln_id=vuln_id,
                title=title,
                description=description or "N/A",
                risk_level=risk_level,
                affected_components=affected or "N/A",
                recommendation=recommendation or "N/A",
            )

        except Exception as e:
            logger.debug(f"Could not build vulnerability from data: {e}")
            return None

    def _find_field(self, data: Dict[str, str], keys: List[str]) -> str:
        """
        Find field in data dictionary using multiple possible keys.

        Args:
            data: Dictionary with lowercase keys
            keys: List of possible key names to search for

        Returns:
            Found value or empty string
        """
        for key in keys:
            if key in data:
                return data[key]
        return ""

    def _parse_cvss_score(self, score_str: str) -> Optional[float]:
        """
        Parse CVSS score from string.

        Args:
            score_str: String containing CVSS score

        Returns:
            Float score or None
        """
        if not score_str:
            return None

        try:
            # Extract first number found
            match = re.search(r"\d+\.?\d*", score_str)
            if match:
                score = float(match.group())
                if 0.0 <= score <= 10.0:
                    return score
        except Exception:
            pass

        return None

    def _map_risk_level(self, risk_text: str) -> RiskLevel:
        """
        Map risk level text to RiskLevel enum.

        Args:
            risk_text: Text describing risk level

        Returns:
            RiskLevel enum value
        """
        risk_lower = risk_text.lower()

        if "critical" in risk_lower or "crit" in risk_lower:
            return RiskLevel.CRITICAL
        elif "high" in risk_lower:
            return RiskLevel.HIGH
        elif "medium" in risk_lower or "med" in risk_lower:
            return RiskLevel.MEDIUM
        elif "low" in risk_lower:
            return RiskLevel.LOW
        else:
            return RiskLevel.INFORMATIONAL

    def _infer_risk_from_id(self, vuln_id: str) -> RiskLevel:
        """
        Infer risk level from vulnerability ID prefix.

        Args:
            vuln_id: Vulnerability ID (e.g., H1, M2)

        Returns:
            RiskLevel enum value
        """
        prefix = vuln_id[0].upper() if vuln_id else "I"

        risk_map = {
            "C": RiskLevel.CRITICAL,
            "H": RiskLevel.HIGH,
            "M": RiskLevel.MEDIUM,
            "L": RiskLevel.LOW,
            "I": RiskLevel.INFORMATIONAL,
        }

        return risk_map.get(prefix, RiskLevel.INFORMATIONAL)

    def _deduplicate_vulnerabilities(
        self, vulnerabilities: List[Vulnerability]
    ) -> List[Vulnerability]:
        """
        Remove duplicate vulnerabilities based on vuln_id.

        Args:
            vulnerabilities: List of vulnerabilities

        Returns:
            Deduplicated list
        """
        seen_ids = set()
        unique = []

        for vuln in vulnerabilities:
            if vuln.vuln_id not in seen_ids:
                seen_ids.add(vuln.vuln_id)
                unique.append(vuln)

        return unique
