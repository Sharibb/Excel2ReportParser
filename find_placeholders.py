"""Find all placeholders in Word template."""

from docx import Document
import re

doc = Document('WAPT-Rootnik-Technical.docx')

print("="*80)
print("FINDING ALL PLACEHOLDERS IN TEMPLATE")
print("="*80)

placeholders = set()

# Check tables
for table_idx, table in enumerate(doc.tables):
    for row in table.rows:
        for cell in row.cells:
            # Find all {{...}} patterns
            found = re.findall(r'\{\{[^}]+\}\}', cell.text)
            placeholders.update(found)

# Check paragraphs  
for para in doc.paragraphs:
    found = re.findall(r'\{\{[^}]+\}\}', para.text)
    placeholders.update(found)

print(f"\nFound {len(placeholders)} unique placeholders:\n")
for p in sorted(placeholders):
    print(f"  - {p}")

print("\n" + "="*80)

# Now specifically check for POC-related content
print("\nSEARCHING FOR 'POC' OR 'STEPS' TEXT (case-insensitive):")
print("="*80)

poc_mentions = 0
for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            text_lower = cell.text.lower()
            if 'poc' in text_lower or 'proof of concept' in text_lower or 'steps to reproduce' in text_lower:
                poc_mentions += 1
                print(f"\n[Table {table_idx+1}, Row {row_idx+1}, Cell {cell_idx+1}]")
                print(f"  {cell.text[:200]}")

print(f"\n{'='*80}")
print(f"Total cells mentioning PoC/Steps: {poc_mentions}")
