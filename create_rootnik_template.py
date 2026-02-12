"""
Template Generator matching WAPT-Rootnik-Technical.docx structure
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_template():
    doc = Document()
    
    # ========================================================================
    # TITLE SECTION
    # ========================================================================
    
    # Annexure heading
    title = doc.add_paragraph()
    run = title.add_run('Annexure I')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 176, 240)  # Blue color (00B0F0)
    
    # Web App URL
    heading1 = doc.add_heading('Web App URL: {{APP_URL}}', level=1)
    for run in heading1.runs:
        run.font.color.rgb = RGBColor(0, 176, 240)
    
    doc.add_paragraph()  # Empty line
    
    # Description paragraph
    desc = doc.add_paragraph()
    desc_run = desc.add_run(
        'The table below summarizes the findings for OWASP Top 10 list for web application security risks.'
    )
    desc_run.font.size = Pt(11)
    desc.style = 'Body Text'
    
    doc.add_paragraph()  # Empty line
    
    # Recommendation paragraph
    rec = doc.add_paragraph()
    rec_run = rec.add_run(
        'RootNik Labs recommends that the Client address the findings outlined in this report. '
        'The findings are categorized based on their severity level.'
    )
    rec_run.font.size = Pt(11)
    rec.style = 'Body Text'
    
    doc.add_paragraph()  # Empty line
    doc.add_page_break()
    
    # ========================================================================
    # ASSESSMENT FINDINGS SECTION
    # ========================================================================
    
    findings_heading = doc.add_heading('Assessment Findings:', level=1)
    
    doc.add_paragraph()  # Empty line
    
    # Summary sub-heading
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run('Summary:')
    summary_run.bold = True
    summary_run.font.size = Pt(18)
    summary_run.font.color.rgb = RGBColor(0, 176, 240)
    summary_para.style = 'Body Text'
    
    # Summary description
    summary_desc = doc.add_paragraph()
    summary_desc_run = summary_desc.add_run(
        'The table below outlines a summary of findings identified during the assessment:'
    )
    summary_desc_run.font.size = Pt(11)
    summary_desc.style = 'Body Text'
    
    doc.add_paragraph()  # Empty line
    
    # Summary table
    summary_table = doc.add_table(rows=8, cols=2)
    summary_table.style = 'Normal Table'
    
    summary_table.rows[0].cells[0].text = 'Finding Description'
    summary_table.rows[0].cells[1].text = 'Status'
    
    summary_table.rows[1].cells[0].text = 'Critical Risk Findings'
    summary_table.rows[1].cells[1].text = 'NONE'
    
    # Empty row
    summary_table.rows[2].cells[0].text = ''
    summary_table.rows[2].cells[1].text = ''
    
    summary_table.rows[3].cells[0].text = 'High Risk Findings'
    summary_table.rows[4].cells[0].text = '{{HIGH_FINDINGS_LIST}}'
    summary_table.rows[4].cells[1].text = 'HIGH'
    
    # Empty row
    summary_table.rows[5].cells[0].text = ''
    summary_table.rows[5].cells[1].text = ''
    
    summary_table.rows[6].cells[0].text = 'Low Risk Findings'
    summary_table.rows[7].cells[0].text = '{{LOW_FINDINGS_LIST}}'
    summary_table.rows[7].cells[1].text = 'LOW'
    
    doc.add_paragraph()  # Empty line
    doc.add_page_break()
    
    # ========================================================================
    # CRITICAL RISK FINDINGS SECTION
    # ========================================================================
    
    critical_heading = doc.add_heading('Critical Risk Findings', level=1)
    for run in critical_heading.runs:
        run.font.color.rgb = RGBColor(0, 176, 240)
        run.font.size = Pt(18)
    
    critical_desc = doc.add_paragraph()
    critical_desc_run = critical_desc.add_run('RootNik Labs found no critical-severity issues.')
    critical_desc_run.font.size = Pt(11)
    critical_desc.style = 'Body Text'
    
    doc.add_paragraph()  # Empty line
    doc.add_page_break()
    
    # ========================================================================
    # HIGH RISK FINDINGS SECTION (TEMPLATE)
    # ========================================================================
    
    high_risk_heading = doc.add_heading('High Risk Findings', level=1)
    for run in high_risk_heading.runs:
        run.font.color.rgb = RGBColor(0, 176, 240)
        run.font.size = Pt(18)
    
    high_desc = doc.add_paragraph()
    high_desc_run = high_desc.add_run(
        'RootNik Labs found {{HIGH_COUNT}} high-severity issues, as described below:'
    )
    high_desc_run.font.size = Pt(11)
    high_desc.style = 'Body Text'
    
    doc.add_paragraph()  # Empty line
    
    # Vulnerability title (This will be duplicated for each vulnerability)
    vuln_title = doc.add_heading('{{VULN_ID}}. {{TITLE}}', level=3)
    for run in vuln_title.runs:
        run.font.color.rgb = RGBColor(238, 0, 0)  # Red color (EE0000)
        run.font.size = Pt(12)
    
    doc.add_paragraph()  # Empty line
    
    # ========================================================================
    # VULNERABILITY DETAILS TABLE (TEMPLATE TO BE DUPLICATED)
    # ========================================================================
    
    vuln_table = doc.add_table(rows=10, cols=2)
    vuln_table.style = 'Normal Table'
    
    # Define table rows
    table_rows = [
        ('Severity:', '{{RISK_LEVEL}}'),
        ('Remediation Efforts:', '{{REMEDIATION_EFFORT}}'),
        ('CVSS:', '{{CVSS_SCORE}}'),
        ('CVE / CWE ID:', '{{CWE_ID}}'),
        ('Summary:', '{{DESCRIPTION}}'),
        ('Affected Assets / Parameters:', '{{AFFECTED_COMPONENTS}}'),
        ('Steps to Reproduce:', '{{POC}}'),
        ('Impact:', '{{IMPACT}}'),
        ('Recommendations:', '{{RECOMMENDATION}}'),
        ('References:', '{{REFERENCES}}'),
    ]
    
    for idx, (label, placeholder) in enumerate(table_rows):
        row = vuln_table.rows[idx]
        
        # Set label (bold)
        label_cell = row.cells[0]
        label_para = label_cell.paragraphs[0]
        label_run = label_para.add_run(label)
        label_run.bold = True
        label_cell.text = ''  # Clear default text
        label_para.add_run(label)
        label_para.runs[0].bold = True
        
        # Set placeholder value
        value_cell = row.cells[1]
        value_cell.text = placeholder
    
    doc.add_paragraph()  # Empty line after table
    doc.add_page_break()
    
    # ========================================================================
    # LOW RISK FINDINGS SECTION (TEMPLATE)
    # ========================================================================
    
    low_risk_heading = doc.add_heading('Low Risk Findings', level=1)
    for run in low_risk_heading.runs:
        run.font.color.rgb = RGBColor(0, 176, 240)
        run.font.size = Pt(18)
    
    low_desc = doc.add_paragraph()
    low_desc_run = low_desc.add_run(
        'RootNik Labs found {{LOW_COUNT}} low-severity issues, as described below:'
    )
    low_desc_run.font.size = Pt(11)
    low_desc.style = 'Body Text'
    
    doc.add_paragraph()  # Empty line
    
    # Vulnerability title for low risk (This will be duplicated for each vulnerability)
    vuln_title_low = doc.add_heading('{{VULN_ID}}. {{TITLE}}', level=3)
    for run in vuln_title_low.runs:
        run.font.color.rgb = RGBColor(238, 0, 0)  # Red color (EE0000)
        run.font.size = Pt(12)
    
    doc.add_paragraph()  # Empty line
    
    # Low risk vulnerability table (same structure as high risk)
    vuln_table_low = doc.add_table(rows=10, cols=2)
    vuln_table_low.style = 'Normal Table'
    
    for idx, (label, placeholder) in enumerate(table_rows):
        row = vuln_table_low.rows[idx]
        
        # Set label (bold)
        label_cell = row.cells[0]
        label_para = label_cell.paragraphs[0]
        label_cell.text = ''
        label_run = label_para.add_run(label)
        label_run.bold = True
        
        # Set placeholder value
        value_cell = row.cells[1]
        value_cell.text = placeholder
    
    doc.add_paragraph()  # Empty line after table
    
    # ========================================================================
    # SAVE TEMPLATE
    # ========================================================================
    
    doc.save('/app/Vulnerability_Report_Template_RootNik.docx')
    print('✅ RootNik-style template created successfully!')
    print('   Template: Vulnerability_Report_Template_RootNik.docx')
    print()
    print('📋 Template includes:')
    print('   - Annexure title section')
    print('   - Assessment findings summary')
    print('   - Critical/High/Low risk sections')
    print('   - Vulnerability detail tables with 10 fields')
    print('   - Proper formatting (colors, fonts, styles)')
    print()
    print('🎨 Formatting:')
    print('   - Blue headings (RGB 00B0F0)')
    print('   - Red vulnerability titles (RGB EE0000)')
    print('   - 11pt body text')
    print('   - 18pt section headings')
    print('   - 12pt vulnerability titles')

if __name__ == '__main__':
    create_template()
