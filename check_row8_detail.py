"""Check row 8 (Recommendations) in detail"""
from docx import Document

doc = Document("/app/test_template.docx")

print("=" * 80)
print("DETAILED CHECK OF RECOMMENDATIONS ROW")
print("=" * 80)

# Find first vulnerability table
for table_idx, table in enumerate(doc.tables):
    table_text = ""
    for row in table.rows:
        for cell in row.cells:
            table_text += cell.text
    
    if '{{RISK' in table_text and '{{CVSS' in table_text:
        print(f"\nChecking table #{table_idx}")
        print("-" * 80)
        
        if len(table.rows) > 8:
            row = table.rows[8]
            if len(row.cells) >= 2:
                label_cell = row.cells[0]
                value_cell = row.cells[1]
                
                print(f"Label: {label_cell.text}")
                print(f"\nValue cell FULL TEXT:")
                print(f"  '{value_cell.text}'")
                print(f"\nChecking for placeholders:")
                print(f"  Contains '{{{{RECOMMENDATION}}}}': {'{{RECOMMENDATION}}' in value_cell.text}")
                print(f"  Contains '{{{{REMEDIATION}}}}': {'{{REMEDIATION}}' in value_cell.text}")
                
                print(f"\nParagraphs in value cell:")
                for p_idx, para in enumerate(value_cell.paragraphs):
                    print(f"  Paragraph {p_idx}: '{para.text}'")
        
        break  # Only check first table

print("\n" + "=" * 80)
