"""Detailed inspection of generated document tables"""
from docx import Document

doc = Document("generated_for_inspection.docx")

print("=" * 80)
print(f"TOTAL TABLES IN DOCUMENT: {len(doc.tables)}")
print("=" * 80)

for table_idx, table in enumerate(doc.tables):
    print(f"\n{'='*80}")
    print(f"TABLE #{table_idx}")
    print(f"{'='*80}")
    print(f"Rows: {len(table.rows)}, Columns: {len(table.columns)}")
    print(f"\nFirst 3 rows:")
    for row_idx, row in enumerate(table.rows[:3]):
        cells_text = [cell.text[:40] + "..." if len(cell.text) > 40 else cell.text for cell in row.cells]
        print(f"  Row {row_idx}: {' | '.join(cells_text)}")
    
    # Check if this looks like a vulnerability table
    if len(table.rows) > 5:
        print(f"\n  → Looks like a vulnerability table (has {len(table.rows)} rows)")
        # Show more details
        print(f"\n  Sample data from rows:")
        for row_idx in [0, 1, 2, 5, 8]:
            if row_idx < len(table.rows):
                row = table.rows[row_idx]
                if len(row.cells) >= 2:
                    cell0 = row.cells[0].text[:30]
                    cell1 = row.cells[1].text[:80]
                    print(f"    Row {row_idx}: {cell0:30} | {cell1}")

print("\n" + "=" * 80)
