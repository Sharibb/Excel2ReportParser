"""Diagnose summary table placeholder replacement issue"""
from pathlib import Path
from app.services.phase2.excel_reader import ExcelReader
from docx import Document
from docx.oxml.ns import qn

# Read Excel data
excel_path = Path("All_Risk_Levels_Template.xlsx")
reader = ExcelReader(excel_path)
report = reader.read()

print("=" * 80)
print("EXCEL DATA ANALYSIS")
print("=" * 80)
print(f"Total vulnerabilities: {len(report.vulnerabilities)}\n")

# Group by risk level
from collections import defaultdict
vulns_by_risk = defaultdict(list)
for v in report.vulnerabilities:
    vulns_by_risk[v.risk_level].append(v)

for risk_level, vulns in sorted(vulns_by_risk.items()):
    print(f"{risk_level}: {len(vulns)} vulnerabilities")
    for v in vulns:
        print(f"  - {v.vuln_id}: {v.title}")

print(f"\nReport counts:")
print(f"  Critical: {report.critical_count}")
print(f"  High: {report.high_count}")
print(f"  Medium: {report.medium_count}")
print(f"  Low: {report.low_count}")
print(f"  Info: {report.info_count}")

# Find and analyze template tables
print("\n" + "=" * 80)
print("TEMPLATE ANALYSIS")
print("=" * 80)

template_path = Path("All_Risk_Levels_Template.docx")
if template_path.exists():
    doc = Document(template_path)
    
    print(f"Found {len(doc.tables)} tables in template\n")
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTable {table_idx}:")
        print(f"  Rows: {len(table.rows)}, Cols: {len(table.columns)}")
        
        # Check for placeholders
        placeholders = []
        for row_idx, row in enumerate(table.rows):
            row_text = " ".join([cell.text for cell in row.cells])
            
            # Check for summary table placeholders
            if any(p in row_text for p in [
                "{{CRITICAL_FINDINGS_LIST}}", "{{HIGH_FINDINGS_LIST}}", 
                "{{MEDIUM_FINDINGS_LIST}}", "{{LOW_FINDINGS_LIST}}", 
                "{{INFO_FINDINGS_LIST}}"
            ]):
                placeholder = None
                for p in ["{{CRITICAL_FINDINGS_LIST}}", "{{HIGH_FINDINGS_LIST}}", 
                         "{{MEDIUM_FINDINGS_LIST}}", "{{LOW_FINDINGS_LIST}}", 
                         "{{INFO_FINDINGS_LIST}}"]:
                    if p in row_text:
                        placeholder = p
                        break
                
                print(f"  Row {row_idx}: Found placeholder '{placeholder}'")
                print(f"    Row text: '{row_text[:150]}'")
                
                # Check if placeholder is split across text elements
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text
                    if placeholder and placeholder in cell_text:
                        print(f"    Cell {cell_idx} contains placeholder")
                        
                        # Check XML structure
                        for para in cell.paragraphs:
                            text_elements = list(para._element.iter(qn('w:t')))
                            if len(text_elements) > 1:
                                print(f"      Paragraph has {len(text_elements)} text elements:")
                                for t_idx, t_elem in enumerate(text_elements):
                                    print(f"        [{t_idx}]: '{t_elem.text}'")
                            else:
                                print(f"      Paragraph has 1 text element: '{text_elements[0].text if text_elements else 'NONE'}'")
else:
    print("Template file not found!")

print("\n" + "=" * 80)
print("DIAGNOSIS COMPLETE")
print("=" * 80)
