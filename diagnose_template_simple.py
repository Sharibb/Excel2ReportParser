"""Simple template analysis without app dependencies"""
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
import openpyxl

# Find template file
template_path = Path("All_Risk_Levels_Template.docx")
excel_path = Path("All_Risk_Levels_Template.xlsx")

print("=" * 80)
print("EXCEL DATA ANALYSIS")
print("=" * 80)

if excel_path.exists():
    wb = openpyxl.load_workbook(excel_path)
    ws = wb.active
    
    print(f"Excel has {ws.max_row} rows\n")
    
    # Read header
    headers = [cell.value for cell in ws[1]]
    print("Headers:", headers[:10])
    
    # Count by risk level
    risk_levels = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0]:  # Has vuln ID
            risk = str(row[headers.index('Risk Level')] if 'Risk Level' in headers else row[3])
            risk_levels[risk] = risk_levels.get(risk, 0) + 1
    
    print("\nVulnerabilities by risk level:")
    for risk, count in sorted(risk_levels.items()):
        print(f"  {risk}: {count}")
else:
    print("Excel file not found!")

print("\n" + "=" * 80)
print("TEMPLATE ANALYSIS")
print("=" * 80)

if template_path.exists():
    doc = Document(template_path)
    
    print(f"\nFound {len(doc.tables)} tables in template\n")
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTable {table_idx} ({len(table.rows)} rows x {len(table.columns)} cols):")
        
        # Show first few rows
        for row_idx in range(min(5, len(table.rows))):
            row = table.rows[row_idx]
            row_text = " | ".join([cell.text[:50] for cell in row.cells])
            print(f"  Row {row_idx}: {row_text}")
            
            # Check for findings list placeholders
            for cell in row.cells:
                cell_text = cell.text
                if "FINDINGS_LIST" in cell_text:
                    print(f"    ⚠️  Found FINDINGS_LIST placeholder!")
                    
                    # Analyze XML structure
                    for para_idx, para in enumerate(cell.paragraphs):
                        text_elements = list(para._element.iter(qn('w:t')))
                        print(f"      Paragraph {para_idx}: {len(text_elements)} text elements")
                        
                        if len(text_elements) > 1:
                            print(f"        ⚠️  SPLIT PLACEHOLDER DETECTED!")
                            full_text = "".join([t.text or "" for t in text_elements])
                            print(f"        Full text: '{full_text}'")
                            for t_idx, t_elem in enumerate(text_elements):
                                print(f"          Element [{t_idx}]: '{t_elem.text}'")
                        else:
                            if text_elements:
                                print(f"        Text: '{text_elements[0].text}'")
else:
    print("Template file not found!")

print("\n" + "=" * 80)
print("ROOT CAUSE IDENTIFIED")
print("=" * 80)
