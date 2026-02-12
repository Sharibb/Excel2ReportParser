"""Diagnose data mismatch between Excel and generated Word document"""
import sys
sys.path.insert(0, '/app')

from pathlib import Path
from app.services.phase2.excel_reader import ExcelReader
from docx import Document

print("=" * 80)
print("DIAGNOSING DATA MISMATCH")
print("=" * 80)

# Read Excel
excel_path = Path("/app/test.xlsx")
print(f"\n1. READING EXCEL: {excel_path}")
print("-" * 80)

reader = ExcelReader(excel_path)
report = reader.read()

print(f"Total vulnerabilities read: {len(report.vulnerabilities)}\n")

for idx, vuln in enumerate(report.vulnerabilities, 1):
    print(f"{idx}. {vuln.vuln_id}: {vuln.title}")
    print(f"   Risk Level: {vuln.risk_level}")
    print(f"   CVSS: {vuln.cvss_score}")
    print(f"   CWE: {vuln.cwe_id}")
    print(f"   Description (first 100 chars): {vuln.description[:100] if vuln.description else 'None'}...")
    print(f"   Affected Components: {vuln.affected_components[:100] if vuln.affected_components else 'None'}...")
    print(f"   Impact (first 100 chars): {vuln.impact[:100] if vuln.impact else 'None'}...")
    print(f"   Recommendation (first 100 chars): {vuln.recommendation[:100] if vuln.recommendation else 'None'}...")
    print(f"   References: {vuln.references[:100] if vuln.references else 'None'}...")
    print(f"   Steps: {len(vuln.steps)} steps")
    if vuln.steps:
        for step_idx, step in enumerate(vuln.steps[:2], 1):  # Show first 2 steps
            print(f"      Step {step_idx}: {step[:80]}...")
    print()

# Check generated document
doc_path = Path("/app/test_output.docx")
if doc_path.exists():
    print("\n2. CHECKING GENERATED DOCUMENT")
    print("-" * 80)
    doc = Document(doc_path)
    
    # Count tables
    print(f"Total tables in document: {len(doc.tables)}")
    
    # Check for placeholders
    print("\n3. CHECKING FOR REMAINING PLACEHOLDERS")
    print("-" * 80)
    placeholders_found = []
    
    for para in doc.paragraphs:
        if '{{' in para.text:
            placeholders_found.append(para.text[:100])
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '{{' in cell.text:
                    placeholders_found.append(cell.text[:100])
    
    if placeholders_found:
        print(f"⚠️  Found {len(placeholders_found)} placeholders:")
        for p in placeholders_found[:10]:  # Show first 10
            print(f"   - {p}")
    else:
        print("✅ No placeholders found - all replaced!")
    
    # Check first vulnerability table (skip summary table at index 0)
    print("\n4. CHECKING FIRST VULNERABILITY TABLE (Table #1)")
    print("-" * 80)
    if len(doc.tables) > 1:
        table = doc.tables[1]
        print(f"Table has {len(table.rows)} rows, {len(table.columns)} columns")
        print("\nFirst 5 rows:")
        for row_idx, row in enumerate(table.rows[:5]):
            print(f"  Row {row_idx}: {' | '.join([cell.text[:50] for cell in row.cells[:3]])}")

print("\n" + "=" * 80)
